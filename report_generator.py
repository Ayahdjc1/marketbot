import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import requests
import json
from datetime import datetime
from database import get_engagement_data_by_range
from config import OLLAMA_API_URL
import logging
from database import get_engagement_data

logger = logging.getLogger(__name__)

def hex_to_rgb_color(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip('#')
    r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    return RGBColor(r, g, b)

class ReportGenerator:
    def __init__(self):
        self.ollama_url = OLLAMA_API_URL
        self.headers = {"Content-Type": "application/json"}
        self._validate_connection()
        self._init_styles()

    def _init_styles(self):
        self.styles = {
            'title': {'font_size': 16, 'bold': True, 'color': '#1F497D'},
            'header1': {'font_size': 14, 'color': '#2E74B5'},
            'header2': {'font_size': 12, 'color': '#2E75B6'},
            'body': {'font_size': 11}
        }

    def _apply_style(self, paragraph, style_name):
        style = self.styles[style_name]
        run = paragraph.runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(style['font_size'])
        if style.get('bold'):
            run.bold = True
        if style.get('color'):
            run.font.color.rgb = hex_to_rgb_color(style['color'])

    def _validate_connection(self):
        try:
            requests.get(self.ollama_url, timeout=5)
        except Exception as e:
            logger.error(f"Ошибка подключения к Ollama: {str(e)}")
            raise RuntimeError("Сервер Ollama недоступен")

    def _query_llama(self, prompt: str) -> str:
        data = {
            "model": "llama3",
            "prompt": f"{prompt}\nОтвет должен быть на русском языке.",
            "stream": False,
            "options": {"temperature": 0.5, "top_p": 0.9}
        }
        try:
            response = requests.post(self.ollama_url, headers=self.headers, data=json.dumps(data), timeout=30)
            response.raise_for_status()
            return response.json().get("response", "Не удалось получить ответ")
        except requests.exceptions.RequestException as e:
            logger.error(f"Ошибка запроса: {str(e)}")
            return "Ошибка анализа данных"

    def _prepare_dataframe(self, raw_data: list, columns: list) -> pd.DataFrame:
        df = pd.DataFrame(raw_data, columns=columns)
        if df.empty:
            raise ValueError("Нет данных для анализа")
        df['date'] = pd.to_datetime(df['date'], errors='coerce')
        if df['date'].isnull().any():
            raise ValueError("Некорректный формат даты в данных")
        return df.sort_values('date').reset_index(drop=True)

    def _create_chart(self, fig_func) -> io.BytesIO:
        img_stream = io.BytesIO()
        fig_func()
        plt.savefig(img_stream, format='png', bbox_inches='tight', dpi=120)
        img_stream.seek(0)
        plt.close()
        return img_stream

    def _create_metric_comparison(self, df):
        def plot():
            metrics = df[['likes', 'comments', 'shares']].mean()
            sns.barplot(x=metrics.index, y=metrics.values, palette="viridis")
            plt.title("Сравнение средних показателей")
            plt.ylabel("Количество")
        return self._create_chart(plot)

    def _create_correlation_matrix(self, df):
        def plot():
            corr = df[['likes', 'comments', 'shares']].corr()
            sns.heatmap(corr, annot=True, cmap='coolwarm', fmt=".2f")
            plt.title("Корреляция между показателями")
        return self._create_chart(plot)

    def _create_weekly_trend(self, df):
        def plot():
            df['weekday'] = df['date'].dt.day_name()
            sns.lineplot(data=df, x='weekday', y='likes', estimator='sum', sort=False, marker='o')
            plt.title("Активность по дням недели")
        return self._create_chart(plot)

    def _add_table(self, doc, data, headers):
        table = doc.add_table(rows=1, cols=len(headers), style='Table Grid')
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(header)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for row in data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                cell = row_cells[i]
                cell.text = str(value)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _generate_chart_analysis_individual(self, chart_id: int, df: pd.DataFrame) -> str:
        prompt = f"Проанализируй график {chart_id} на основе данных по лайкам, комментариям и репостам. Выдели ключевые тренды."
        return self._query_llama(prompt)

    def _generate_chart_analysis(self, df: pd.DataFrame) -> str:
        metrics = df[['likes', 'comments', 'shares']].mean()
        corr = df[['likes', 'comments', 'shares']].corr()
        df['weekday'] = df['date'].dt.day_name()
        weekday_likes = df.groupby('weekday')['likes'].sum()

        summary = f"""
График 1. Средние значения: лайки {metrics['likes']:.2f}, комментарии {metrics['comments']:.2f}, репосты {metrics['shares']:.2f}.
График 2. Корреляции: лайки-комментарии {corr.loc['likes', 'comments']:.2f}, лайки-репосты {corr.loc['likes', 'shares']:.2f}, комментарии-репосты {corr.loc['comments', 'shares']:.2f}.
График 3. Активность по дням недели:
{weekday_likes.to_string()}
"""
        return self._query_llama("Проанализируй эти графики на основе данных. Сделай выводы.\n" + summary)

    def _generate_advanced_analysis(self, df):
        prompt = f"""
Проанализируй данные с {df['date'].min().date()} по {df['date'].max().date()}:
- выяви сезонные паттерны активности
- определи посты с аномальной вовлеченностью
- сравни эффективность разных типов контента
- предложи лучшее время публикаций
"""
        return self._query_llama(prompt)
    def generate_report(self, period: str) -> str:
        raw_data = get_engagement_data(period)
        df = self._prepare_dataframe(raw_data, ['id', 'post_id', 'likes', 'comments', 'shares', 'date', 'channel'])
        doc = self._create_report_document(df, f"Отчет о вовлеченности за период: {period}")
        filename = f"Отчет_период_{period}_{datetime.now().strftime('%Y-%m-%d')}.docx"
        doc.save(filename)
        return filename
    def _generate_recommendations(self, df):
        prompt = f"""
На основе данных по лайкам, комментариям и репостам с {df['date'].min().date()} по {df['date'].max().date()}, дай рекомендации:
- по времени публикаций
- по формату и темам
- по улучшению взаимодействия
"""
        return self._query_llama(prompt)

    def _create_report_document(self, df, title):
        doc = Document()
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._apply_style(title_para, 'title')

        doc.add_heading("Содержание", level=1)
        doc.add_paragraph("1. Ключевые показатели\n2. Визуальная аналитика\n3. Анализ визуализаций\n4. Глубинный анализ\n5. Рекомендации\n6. Заключение")

        doc.add_heading("1. Ключевые показатели", level=1)
        stats = df[['likes', 'comments', 'shares']].describe().round(2)
        self._add_table(doc, stats.values.tolist(), stats.columns.tolist())

        doc.add_heading("2. Визуальная аналитика", level=1)
        visual_funcs = [
            ("Сравнение метрик", self._create_metric_comparison(df), 1),
            ("Корреляция показателей", self._create_correlation_matrix(df), 2),
            ("Активность по дням недели", self._create_weekly_trend(df), 3)
        ]
        for title, stream, chart_id in visual_funcs:
            doc.add_heading(title, level=2)
            doc.add_picture(stream, width=Inches(5.5))
            mini_analysis = self._generate_chart_analysis_individual(chart_id, df)
            doc.add_paragraph(mini_analysis)

        doc.add_heading("3. Анализ визуализаций", level=1)
        doc.add_paragraph(self._generate_chart_analysis(df))

        doc.add_heading("4. Глубинный анализ", level=1)
        doc.add_paragraph(self._generate_advanced_analysis(df))

        doc.add_heading("5. Рекомендации", level=1)
        doc.add_paragraph(self._generate_recommendations(df))

        doc.add_heading("6. Заключение", level=1)
        doc.add_paragraph(f"Отчет сгенерирован {datetime.now().strftime('%d.%m.%Y %H:%M')} системой аналитики Telegram-каналов.")

        return doc

    def generate_report_by_date_range(self, start_date: str, end_date: str) -> str:
        try:
            raw_data = get_engagement_data_by_range(start_date, end_date)
            df = self._prepare_dataframe(raw_data, ['id', 'post_id', 'likes', 'comments', 'shares', 'date', 'channel'])
            doc = self._create_report_document(df, f"Отчет о вовлеченности {start_date} - {end_date}")
            filename = f"Отчет_{start_date}_по_{end_date}.docx"
            doc.save(filename)
            return filename
        except Exception as e:
            logger.error(f"Ошибка генерации отчета: {str(e)}")
            raise

if __name__ == "__main__":
    generator = ReportGenerator()
    try:
        path = generator.generate_report_by_date_range("2024-01-01", "2024-01-07")
        print(f"Отчет сохранен: {path}")
    except Exception as err:
        print(f"Ошибка: {err}")

